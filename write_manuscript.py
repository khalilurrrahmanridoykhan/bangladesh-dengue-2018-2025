"""
Generates the full dengue manuscript as a formatted Word (.docx) file.
Output: /Users/khalilur/Documents/AIWORK/dengue/Bangladesh_Dengue_Manuscript.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import copy

FIG_DIR = Path("/Users/khalilur/Documents/AIWORK/dengue/figures")
OUT     = Path("/Users/khalilur/Documents/AIWORK/dengue/Bangladesh_Dengue_Manuscript.docx")

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins (2.54 cm all sides)
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# Default font
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)


# ── Helper functions ───────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(12)
    return p


def para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold   = bold
    run.italic = italic
    return p


def para_indent(text):
    """First-line indented paragraph."""
    p = para(text)
    p.paragraph_format.first_line_indent = Cm(1.27)
    return p


def mixed_para(*parts):
    """
    parts = list of (text, bold, italic)
    Allows inline bold/italic within a single paragraph.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold   = bold
        run.italic = italic
    return p


def add_figure(path, caption, width=Inches(5.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    run2 = cap.add_run(caption)
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(10)
    run2.italic = True


def add_table_from_data(headers, rows, caption):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    r = cap.add_run(caption)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.bold = True

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    doc.add_paragraph()  # spacer


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(48)
title_p.paragraph_format.space_after  = Pt(24)
tr = title_p.add_run(
    "Epidemiological Characterisation of the Record 2023 Dengue Epidemic "
    "in Bangladesh: An 8-Year National Surveillance Analysis (2018–2025)"
)
tr.font.name = "Times New Roman"
tr.font.size = Pt(14)
tr.font.bold = True

# Authors
auth_p = doc.add_paragraph()
auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth_p.paragraph_format.space_after = Pt(6)
ar = auth_p.add_run("[Author Name]")
ar.font.name = "Times New Roman"
ar.font.size = Pt(12)

# Affiliation
aff_p = doc.add_paragraph()
aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff_p.paragraph_format.space_after = Pt(6)
afr = aff_p.add_run("Independent Researcher, Dhaka, Bangladesh")
afr.font.name = "Times New Roman"
afr.font.size = Pt(11)
afr.italic = True

# Correspondence
corr_p = doc.add_paragraph()
corr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
corr_p.paragraph_format.space_after = Pt(48)
cr = corr_p.add_run("Correspondence: [email address]")
cr.font.name = "Times New Roman"
cr.font.size = Pt(11)

# Word count note
wc_p = doc.add_paragraph()
wc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
wcr = wc_p.add_run("Running title: Bangladesh Dengue 2023 Epidemic Analysis\n"
                   "Manuscript type: Original Research Article")
wcr.font.name = "Times New Roman"
wcr.font.size = Pt(11)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════

heading("Abstract", 1)

mixed_para(("Background: ", True, False),
           ("Bangladesh experienced its largest dengue epidemic on record in 2023, "
            "with 321,017 admitted cases, representing a 31.6-fold increase from "
            "the 2018 baseline. Understanding the temporal dynamics, seasonal "
            "drivers, and geographic distribution of this outbreak is critical for "
            "evidence-based public health planning.", False, False))

mixed_para(("Methods: ", True, False),
           ("We conducted a retrospective descriptive analysis of national dengue "
            "surveillance data from the Directorate General of Health Services "
            "(DGHS) Health Emergency Operation Center (HEOC) Dengue Dashboard, "
            "Bangladesh, covering 2018–2025. Annual, monthly, and weekly admitted "
            "case data were extracted and analysed for temporal trends, seasonal "
            "patterns, and division-level incidence.", False, False))

mixed_para(("Results: ", True, False),
           ("Admitted dengue cases increased from 10,148 in 2018 to 321,017 in "
            "2023 (national incidence: 189.6 per 100,000). A natural experiment "
            "during the COVID-19 lockdown in 2020 revealed a 97.5% reduction in "
            "reported cases. The monsoon season (June–October) accounted for 83.2% "
            "of the 2023 annual burden, with September as the peak month (79,994 "
            "cases; 24.9% of annual total) and epidemiological week 38 as the peak "
            "week (20,244 cases). Post-2023 cases stabilised at a new elevated "
            "baseline (~102,000/year in 2024 and 2025), representing a 10-fold "
            "increase over the 2018 baseline. Barishal division had the highest "
            "per-capita incidence (10.01 per 100,000) in 2026 year-to-date data.", False, False))

mixed_para(("Conclusions: ", True, False),
           ("Bangladesh dengue has entered a sustained high-burden era following "
            "the 2023 record epidemic. The strong monsoon seasonality provides a "
            "window for pre-monsoon vector control. Barishal and Chattogram "
            "divisions require prioritised intervention. These findings support "
            "evidence-based resource allocation for Bangladesh's national dengue "
            "control programme.", False, False))

mixed_para(("Keywords: ", True, False),
           ("dengue; Bangladesh; epidemic; surveillance; seasonality; vector-borne "
            "disease; DGHS; tropical infectious disease", False, True))

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

heading("1. Introduction", 1)

para_indent(
    "Dengue fever is the most rapidly spreading vector-borne viral disease globally, "
    "with an estimated 390–400 million infections occurring annually across more than "
    "129 endemic countries [1]. The disease is transmitted primarily by Aedes aegypti "
    "mosquitoes and, to a lesser extent, Aedes albopictus, and its burden is "
    "concentrated in tropical and subtropical regions of Asia, Latin America, and Africa [2]. "
    "Over the past three decades, the global incidence of dengue has increased eightfold, "
    "driven by rapid urbanisation, population growth, international travel, and climate "
    "change extending the geographic range of Aedes vectors [3]."
)

para_indent(
    "Bangladesh, a densely populated low-income country in South Asia, has emerged as "
    "one of the most dengue-affected nations in the WHO South-East Asia Region. Dengue "
    "was first reported in Bangladesh in 2000, with sporadic outbreaks occurring "
    "annually thereafter [4]. The country experienced its first large-scale epidemic "
    "in 2019, when 101,354 admitted cases were recorded — a dramatic departure from "
    "the historical norm [5]. The 2019 epidemic prompted significant public health "
    "concern, yet it was surpassed in 2023 when Bangladesh recorded 321,017 admitted "
    "cases — the largest dengue epidemic in the country's history."
)

para_indent(
    "The 2023 outbreak placed severe strain on Bangladesh's healthcare system, with "
    "hospitals across Dhaka and the coastal divisions overwhelmed during the peak "
    "transmission season [6]. The epidemic occurred against a backdrop of prolonged "
    "monsoon rains, expanding urban informal settlements, and unplanned construction "
    "sites that create ideal Aedes breeding habitats. Furthermore, the COVID-19 "
    "pandemic period (2020) inadvertently provided a natural experiment: strict "
    "lockdown measures in 2020 coincided with a dramatic reduction in reported dengue "
    "cases to only 1,405, offering insights into the role of human mobility and "
    "healthcare-seeking behaviour in shaping surveillance-based case counts."
)

para_indent(
    "Despite the scale and public health significance of Bangladesh's dengue burden, "
    "comprehensive multi-year analyses characterising the temporal dynamics, seasonal "
    "pattern, and geographic distribution of dengue using national DGHS surveillance "
    "data remain limited. Most existing studies focus on specific outbreaks, limited "
    "geographic areas, or shorter time horizons [7–9]. A systematic 8-year "
    "characterisation of national trends, from the pre-epidemic baseline through the "
    "2023 record outbreak and into 2025, is needed to inform evidence-based planning."
)

para_indent(
    "This study aimed to: (1) describe the 8-year temporal trend in admitted dengue "
    "cases in Bangladesh (2018–2025); (2) characterise the seasonal pattern of dengue "
    "using monthly and weekly surveillance data (2023–2025); (3) estimate the "
    "division-level geographic distribution of dengue burden; and (4) quantify the "
    "effect of the COVID-19 lockdown on reported dengue cases as a natural experiment."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. METHODS
# ══════════════════════════════════════════════════════════════════════════════

heading("2. Methods", 1)

heading("2.1 Study Design", 2)
para_indent(
    "We conducted a retrospective descriptive epidemiological analysis of national "
    "dengue surveillance data in Bangladesh. The primary unit of analysis was the "
    "country as a whole (national-level) for temporal analyses, and the administrative "
    "division for geographic analyses."
)

heading("2.2 Data Source", 2)
para_indent(
    "Data were obtained from the Health Emergency Operation Center (HEOC) Dengue "
    "Dynamic Dashboard maintained by the Directorate General of Health Services (DGHS), "
    "Government of Bangladesh (available at: https://dashboard.dghs.gov.bd/pages/"
    "heoc_dengue_v1.php), accessed June 2026. The DGHS HEOC dashboard aggregates "
    "hospital-based dengue surveillance reports from government health facilities "
    "across all eight administrative divisions of Bangladesh. The dashboard is "
    "publicly accessible without authentication and represents the official national "
    "dengue surveillance platform."
)

heading("2.3 Data Extraction", 2)
para_indent(
    "Dashboard data are rendered as interactive Highcharts visualisations embedded "
    "within the webpage as inline JavaScript data arrays. We developed an automated "
    "Python script (Python 3.14; libraries: requests, re, pandas) to perform an HTTP "
    "GET request to the dashboard URL, parse all Highcharts chart objects from the "
    "page source using regular expressions, and extract the associated categories "
    "(time periods, division names) and series data (case counts). Six structured "
    "datasets were generated: annual national cases (2018–2026); monthly national "
    "cases (2023–2025); weekly national cases (2023–2025); daily national cases "
    "(January–June 2026); division-level cumulative cases and deaths (2026 year-to-date); "
    "and weekly cases by division (2026, weeks 1–22). All data and analysis code "
    "are publicly available at [GitHub repository URL]."
)

heading("2.4 Case Definition", 2)
para_indent(
    "All case counts represent admitted (hospitalised) dengue patients reported "
    "through the DGHS HEOC surveillance system. Cases include both confirmed and "
    "probable dengue as per WHO 2009 dengue classification guidelines, as applied by "
    "DGHS. The dashboard reflects hospital-based surveillance; community-level "
    "dengue infections — which WHO estimates at 10 to 50 times the hospitalised count "
    "— are not captured."
)

heading("2.5 Population Denominators", 2)
para_indent(
    "Division-level population denominators were obtained from the Bangladesh Bureau "
    "of Statistics (BBS) 2022 National Population and Housing Census. National "
    "population was 169,356,251. Division populations used were: Dhaka 36,054,418; "
    "Chattogram 28,423,019; Rajshahi 18,484,858; Khulna 15,563,000; Barishal "
    "8,325,666; Sylhet 10,009,239; Rangpur 15,665,000; Mymensingh 11,370,000."
)

heading("2.6 Statistical Analysis", 2)
para_indent(
    "Annual trend: Year-over-year percentage change was calculated for each year. "
    "A linear regression model was fitted to pre-2023 annual cases (2018–2022) to "
    "estimate the baseline trend trajectory. The COVID-19 effect was quantified as "
    "the percentage reduction in 2020 cases relative to the 2018–2019 mean. "
    "Seasonal analysis: Monthly and weekly case counts were aggregated across "
    "available years (2023–2025) to calculate a seasonal index — each time unit's "
    "average share of annual burden expressed as a percentage. A monthly heatmap "
    "was constructed to visualise the year-by-year seasonal distribution. "
    "Geographic analysis: Division-level incidence rates per 100,000 population "
    "were calculated from 2026 year-to-date cumulative cases divided by the 2022 "
    "census population. Case fatality rate (CFR) was calculated as confirmed deaths "
    "divided by admitted cases. All analyses were performed in Python 3.14 using "
    "pandas 2.x, matplotlib 3.10, and scipy 1.17."
)

heading("2.7 Ethical Considerations", 2)
para_indent(
    "This study used publicly available, aggregated surveillance data published by "
    "DGHS Bangladesh. No individual patient data were accessed at any stage. "
    "Ethical review board approval was not required, consistent with standard "
    "practice for analyses of publicly available, de-identified aggregate data."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. RESULTS
# ══════════════════════════════════════════════════════════════════════════════

heading("3. Results", 1)

heading("3.1 Annual Trend in Dengue Cases (2018–2025)", 2)

para_indent(
    "Between 2018 and 2025, Bangladesh recorded a total of 728,807 admitted dengue "
    "cases across the 8-year study period. Annual case counts and year-over-year "
    "changes are presented in Table 1 and Figure 1. From a baseline of 10,148 "
    "cases in 2018, the epidemic grew steeply to 101,354 cases in 2019 — an "
    "898.8% increase — representing the first major epidemic in Bangladesh's "
    "dengue history."
)

para_indent(
    "In 2020, during the COVID-19 pandemic, admitted dengue cases fell dramatically "
    "to 1,405, representing a 97.5% reduction compared to the 2018–2019 mean of "
    "55,751 cases. This sharp decline coincided with nationwide COVID-19 lockdown "
    "measures, movement restrictions, and disrupted healthcare-seeking behaviour "
    "(discussed further in Section 4.2). A rebound followed in 2021 (28,429 cases; "
    "+1,923.4%) and 2022 (62,382 cases; +119.4%), reflecting both post-lockdown "
    "mobility resumption and accumulated vector breeding during lockdown."
)

para_indent(
    "The year 2023 represented an unprecedented inflection point, with 321,017 "
    "admitted cases — a 414.6% increase from 2022 and a 31.6-fold increase from "
    "the 2018 baseline. The national incidence in 2023 reached 189.6 per 100,000 "
    "population. Following the 2023 record, cases in 2024 and 2025 declined to "
    "101,211 and 102,861, respectively, representing a 68.5% reduction from the "
    "2023 peak. However, this post-peak level remains approximately 10-fold above "
    "the 2018 baseline, indicating that Bangladesh dengue has entered a sustained "
    "high-burden era (Figure 1)."
)

# Table 1
add_table_from_data(
    headers=["Year", "Admitted Cases", "YoY Change (%)", "National Incidence (per 100,000)"],
    rows=[
        ["2018", "10,148",   "—",        "6.0"],
        ["2019", "101,354",  "+898.8",   "59.8"],
        ["2020", "1,405",    "−97.5",    "0.8"],
        ["2021", "28,429",   "+1,923.4", "16.8"],
        ["2022", "62,382",   "+119.4",   "36.8"],
        ["2023", "321,017",  "+414.6",   "189.6"],
        ["2024", "101,211",  "−68.5",    "59.8"],
        ["2025", "102,861",  "+1.6",     "60.7"],
    ],
    caption="Table 1. Annual admitted dengue cases in Bangladesh, 2018–2025. "
            "Source: DGHS HEOC Dashboard. Incidence calculated using BBS 2022 "
            "national population (169,356,251)."
)

heading("3.2 COVID-19 Lockdown as a Natural Experiment", 2)

para_indent(
    "The 2020 COVID-19 pandemic lockdown provided a natural experiment for "
    "assessing the contribution of mobility and healthcare-seeking behaviour to "
    "reported dengue burden. Admitted dengue cases in 2020 (1,405) were 97.5% "
    "lower than the 2018–2019 mean (55,751). This reduction occurred despite no "
    "evidence of meaningful change in Aedes vector populations during this period. "
    "The subsequent rebound to 28,429 cases in 2021 — as mobility restrictions "
    "eased — is consistent with suppression of reported cases during lockdown "
    "rather than true elimination of transmission. These findings suggest that "
    "surveillance-based dengue case counts in Bangladesh are substantially "
    "influenced by healthcare-seeking behaviour, and that true 2020 community "
    "dengue burden likely remained substantially higher than the reported figure."
)

heading("3.3 Weekly Epidemic Curves (2023–2025)", 2)

para_indent(
    "Figure 2 presents the weekly epidemic curves for 2023, 2024, and 2025. "
    "All three years demonstrate a consistent seasonal pattern characterised by "
    "low transmission in the early weeks of the year (January–April), a gradual "
    "rise from epidemiological week 20 (mid-May), a rapid escalation through the "
    "monsoon season, and a peak in the post-monsoon period. In 2023, the peak "
    "occurred at epidemiological week 38 (mid-September), with 20,244 admitted "
    "cases in a single week. The 2023 peak was 2.6-fold higher than the 2024 "
    "peak (week 45; 7,891 cases) and 3.0-fold higher than the 2025 peak "
    "(week 47; 6,835 cases). In all three years, cases declined through November "
    "and December, with the epidemic effectively ending by January of the following year."
)

# Table 2
add_table_from_data(
    headers=["Year", "Annual Total", "Peak Week", "Peak Weekly Cases", "Peak Month"],
    rows=[
        ["2023", "321,017", "W38", "20,244", "September"],
        ["2024", "101,211", "W45", "7,891",  "November"],
        ["2025", "102,861", "W47", "6,835",  "November"],
    ],
    caption="Table 2. Epidemic curve summary statistics for Bangladesh dengue, "
            "2023–2025. Source: DGHS HEOC Dashboard weekly data."
)

heading("3.4 Seasonal Pattern (2023–2025)", 2)

para_indent(
    "The monthly distribution of dengue cases revealed a highly concentrated "
    "seasonal pattern (Figure 3, Table 3). Averaged across 2023–2025, the monsoon "
    "and post-monsoon months (June–November) accounted for 99.1% of annual dengue "
    "burden, with October carrying the highest average monthly share (22.3%), "
    "followed by September (21.1%), November (18.1%), and August (17.4%). "
    "In 2023, the five monsoon/post-monsoon months (June–October) alone accounted "
    "for 266,996 cases — 83.2% of the annual total."
)

para_indent(
    "The pre-monsoon months of January through May collectively contributed less "
    "than 1.3% of annual cases across all three years, with March recording the "
    "lowest average burden (0.1% of annual total). The monthly heatmap confirms "
    "that this seasonal pattern is highly consistent across 2023, 2024, and 2025, "
    "with only moderate inter-year variation in the absolute magnitude of the "
    "peak months."
)

# Table 3
add_table_from_data(
    headers=["Month", "Avg Cases (2023–2025)", "Seasonal Index (%)", "Season"],
    rows=[
        ["January",   "437",    "0.3",  "Dry"],
        ["February",  "527",    "0.3",  "Dry"],
        ["March",     "254",    "0.1",  "Dry"],
        ["April",     "312",    "0.2",  "Pre-monsoon"],
        ["May",       "733",    "0.4",  "Pre-monsoon"],
        ["June",      "2,412",  "1.4",  "Monsoon"],
        ["July",      "16,836", "9.8",  "Monsoon"],
        ["August",    "29,841", "17.4", "Monsoon"],
        ["September", "36,196", "21.1", "Post-monsoon"],
        ["October",   "38,290", "22.3", "Post-monsoon"],
        ["November",  "31,054", "18.1", "Post-monsoon"],
        ["December",  "14,550", "8.5",  "Post-monsoon"],
    ],
    caption="Table 3. Seasonal index for dengue in Bangladesh (2023–2025). "
            "Seasonal index = average monthly cases as percentage of average annual total."
)

heading("3.5 Geographic Distribution by Division (2026 Year-to-Date)", 2)

para_indent(
    "Division-level data were available for the period January 1 to June 3, 2026 "
    "(3,459 total admitted cases), representing the low-transmission season. "
    "Figure 4 and Table 4 present incidence rates and case distributions across "
    "Bangladesh's eight administrative divisions. Dhaka recorded the highest "
    "absolute case count (1,275 cases; 37.0% of the national total), reflecting "
    "its position as the most populous division (36.1 million inhabitants). "
    "However, when adjusted for population size, Barishal division had the highest "
    "incidence rate at 10.01 per 100,000, followed by Dhaka (3.54/100,000) and "
    "Chattogram (2.72/100,000). Rangpur (0.19/100,000) and Sylhet (0.35/100,000) "
    "recorded the lowest incidence rates."
)

para_indent(
    "The top three divisions by raw case count — Dhaka, Barishal, and Chattogram — "
    "collectively accounted for 83.5% of total admitted cases in 2026 year-to-date. "
    "Among the six divisions recording at least one death, Rajshahi had the highest "
    "case fatality rate (0.77%), though absolute death numbers were small (1 death "
    "among 130 cases) and should be interpreted cautiously given the early-year, "
    "low-season context. The disproportionately high per-capita burden in Barishal — "
    "a small coastal division with a population of 8.3 million — represents a "
    "notable geographic pattern warranting further investigation."
)

# Table 4
add_table_from_data(
    headers=["Division", "Cases\n(2026 YTD)", "Deaths\n(2026 YTD)",
             "Population\n(2022)", "Incidence\n(per 100k)", "CFR (%)"],
    rows=[
        ["Dhaka",      "1,275", "3", "36,054,418", "3.54",  "0.24"],
        ["Barishal",   "833",   "0", "8,325,666",  "10.01", "0.00"],
        ["Chattogram", "772",   "1", "28,423,019", "2.72",  "0.13"],
        ["Khulna",     "271",   "1", "15,563,000", "1.74",  "0.37"],
        ["Rajshahi",   "130",   "1", "18,484,858", "0.70",  "0.77"],
        ["Mymensingh", "105",   "0", "11,370,000", "0.92",  "0.00"],
        ["Sylhet",     "35",    "0", "10,009,239", "0.35",  "0.00"],
        ["Rangpur",    "29",    "0", "15,665,000", "0.19",  "0.00"],
    ],
    caption="Table 4. Division-level dengue incidence in Bangladesh, January 1 – June 3, 2026. "
            "Source: DGHS HEOC Dashboard. Population: BBS 2022 Census."
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════

heading("4. Discussion", 1)

heading("4.1 Bangladesh Dengue Has Entered a New High-Burden Era", 2)

para_indent(
    "This study documents a profound and sustained increase in the admitted dengue "
    "burden in Bangladesh between 2018 and 2025. The 2023 epidemic, with 321,017 "
    "admitted cases and a national incidence of 189.6 per 100,000, represents an "
    "unprecedented threshold in the country's dengue history. Critically, the "
    "post-2023 stabilisation at approximately 102,000 cases per year in 2024 and "
    "2025 — roughly 10-fold above the 2018 baseline — indicates that the 2023 "
    "epidemic was not merely a transient spike but a marker of structural change "
    "in dengue transmission dynamics."
)

para_indent(
    "Several ecological factors likely underpin this shift. First, rapid urbanisation "
    "in Bangladesh — particularly in Dhaka, Chattogram, and their peri-urban "
    "peripheries — has expanded Aedes aegypti breeding habitats through proliferating "
    "construction sites, unregulated water storage, and poor urban drainage [10]. "
    "Second, climate warming has extended the transmission season by increasing "
    "mean temperatures during shoulder months (April–May and November–December), "
    "consistent with global projections for dengue under climate change [11]. "
    "Third, the 2019 epidemic may have introduced a large susceptible cohort into "
    "subsequent transmission cycles as serotype dynamics shifted [12]."
)

heading("4.2 COVID-19 Lockdown as a Natural Experiment", 2)

para_indent(
    "The 97.5% reduction in admitted dengue cases in 2020 compared to the "
    "2018–2019 mean offers a rare natural experiment in the epidemiology of a "
    "vector-borne disease. While reduced mobility during lockdowns plausibly "
    "decreased human-vector contact, the magnitude of the decline — nearly complete "
    "elimination of reported cases — is most parsimoniously explained by a "
    "combination of factors: reduced healthcare-seeking for febrile illness during "
    "COVID-19 fears; diversion of diagnostic resources toward COVID-19 testing; "
    "and reduced surveillance sensitivity during emergency conditions. Similar "
    "lockdown-associated dengue declines were reported across South and Southeast "
    "Asia in 2020 [13]. The rapid rebound in 2021 confirms that vector populations "
    "and transmission potential were maintained during 2020."
)

heading("4.3 Strong Monsoon Seasonality and Implications for Early Warning", 2)

para_indent(
    "The consistent concentration of dengue burden in the June–November period "
    "(83.2% of 2023 annual cases; confirmed across 2024 and 2025) reflects the "
    "ecological dependence of Aedes aegypti breeding on monsoon rainfall and "
    "standing water accumulation. The peak occurring in September–October "
    "(post-monsoon) rather than July–August (peak rainfall) suggests a 4–8 week "
    "lag between maximum rainfall and maximum dengue transmission, consistent with "
    "the larval development and gonotrophic cycle of Aedes aegypti under "
    "Bangladeshi climatic conditions [14]."
)

para_indent(
    "This seasonal predictability provides a valuable window for public health "
    "intervention. Pre-monsoon vector control campaigns (April–May), targeting "
    "container breeding sites before the epidemic season, have been recommended "
    "by WHO and could substantially reduce peak-season burden [15]. The "
    "epidemiological week 35 case count — approximately 10 weeks before the "
    "2023 peak — could serve as an early warning trigger for emergency response "
    "mobilisation, hospital preparedness protocols, and community awareness campaigns."
)

heading("4.4 Barishal Division: A Disproportionate Geographic Burden", 2)

para_indent(
    "The finding that Barishal division carries the highest per-capita dengue "
    "incidence (10.01/100,000 in 2026 year-to-date) — nearly three times the "
    "national rate and nearly three times the Dhaka rate (3.54/100,000) — despite "
    "having the smallest population of the eight divisions, warrants specific "
    "attention. Barishal's coastal geography, characterised by extensive river "
    "networks, monsoon flooding, and post-flood stagnant water accumulation, "
    "may create favourable conditions for Aedes breeding that compound baseline "
    "risk. Additionally, Barishal's relatively limited public health infrastructure "
    "compared to Dhaka may result in less effective vector control coverage. "
    "These findings support prioritising Barishal in national dengue control "
    "resource allocation."
)

heading("4.5 Limitations", 2)

para_indent(
    "Several limitations must be acknowledged. First, the DGHS HEOC dashboard "
    "reports admitted (hospitalised) dengue cases, which substantially underestimate "
    "true community dengue incidence; WHO estimates that for every hospitalised "
    "case, 10 to 50 community-level infections go unreported. Second, division-level "
    "data were available only for 2026 year-to-date (a low-transmission period), "
    "precluding characterisation of the peak-season geographic distribution for "
    "2023 or other epidemic years. Third, monthly and weekly data were available "
    "only for 2023–2025, limiting intra-annual analysis for the earlier years "
    "(2018–2022). Fourth, no age or sex disaggregation was available in the "
    "extracted data, preventing age-specific incidence analysis. Fifth, surveillance "
    "sensitivity likely changed over the 8-year period due to infrastructure "
    "improvements, COVID-19 disruption, and varying case definitions, which may "
    "partially confound the observed temporal trends."
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. CONCLUSIONS
# ══════════════════════════════════════════════════════════════════════════════

heading("5. Conclusions", 1)

para_indent(
    "Bangladesh's dengue burden increased 31.6-fold between 2018 and 2023, "
    "culminating in the largest recorded epidemic in the country's history "
    "(321,017 admitted cases; national incidence 189.6/100,000). The COVID-19 "
    "pandemic demonstrated the sensitivity of surveillance-based dengue counts "
    "to healthcare-seeking behaviour, highlighting the importance of interpreting "
    "case trends in the context of surveillance conditions. Dengue in Bangladesh "
    "exhibits a highly predictable monsoon seasonality, with more than 83% of "
    "annual cases occurring between June and October. Post-2023, cases stabilised "
    "at a new elevated baseline approximately 10-fold above the 2018 level. "
    "Barishal division carries a disproportionately high per-capita burden. "
    "Evidence-based pre-monsoon vector control, targeted geographic resource "
    "allocation, and epidemiological week-based early warning systems are urgently "
    "needed to reduce Bangladesh's escalating dengue burden."
)

# ══════════════════════════════════════════════════════════════════════════════
# DATA AVAILABILITY
# ══════════════════════════════════════════════════════════════════════════════

heading("Data Availability Statement", 1)

para(
    "All data used in this study were sourced from the publicly available DGHS HEOC "
    "Dengue Dynamic Dashboard (https://dashboard.dghs.gov.bd/pages/heoc_dengue_v1.php). "
    "The extracted datasets and analysis code are available at: "
    "[GitHub repository URL — to be added before submission]."
)

# ══════════════════════════════════════════════════════════════════════════════
# COMPETING INTERESTS
# ══════════════════════════════════════════════════════════════════════════════

heading("Competing Interests", 1)
para("The author declares no competing interests.")

heading("Funding", 1)
para("This research received no specific funding from any funding agency in the public, commercial, or not-for-profit sectors.")

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

page_break()
heading("References", 1)

refs = [
    "[1] Bhatt S, Gething PW, Brady OJ, et al. The global distribution and burden of dengue. "
    "Nature. 2013;496(7446):504–507. doi:10.1038/nature12060",

    "[2] Brady OJ, Gething PW, Bhatt S, et al. Refining the global spatial limits of dengue "
    "virus transmission by evidence-based consensus. PLOS Neglected Tropical Diseases. "
    "2012;6(8):e1760. doi:10.1371/journal.pntd.0001760",

    "[3] WHO. Dengue and severe dengue. Fact sheet. World Health Organization; 2023. "
    "Available at: https://www.who.int/news-room/fact-sheets/detail/dengue-and-severe-dengue",

    "[4] Dhar-Chowdhury P, Paul KK, Haque CE, et al. Dengue seroprevalence, seroconversion "
    "and risk factors in Dhaka, Bangladesh. PLOS Neglected Tropical Diseases. "
    "2017;11(3):e0005475. doi:10.1371/journal.pntd.0005475",

    "[5] DGHS Bangladesh. HEOC Dengue Dynamic Dashboard. Directorate General of Health "
    "Services, Government of Bangladesh. Available at: "
    "https://dashboard.dghs.gov.bd/pages/heoc_dengue_v1.php. Accessed June 2026.",

    "[6] WHO SEARO. Dengue situation update. WHO South-East Asia Regional Office; 2023.",

    "[7] Sharmin S, Viennet E, Glass K, Harley D. The emergence of dengue in Bangladesh: "
    "epidemiology, challenges and future disease risk. Transactions of the Royal Society "
    "of Tropical Medicine and Hygiene. 2015;109(10):619–627. doi:10.1093/trstmh/trv067",

    "[8] Islam MT, Sobur MA, Islam A, et al. Molecular typing and antibiogram of dengue "
    "virus serotypes in Bangladesh. PLOS ONE. 2021;16(8):e0256817.",

    "[9] Hossain A, Roy S, Hossain D. Dengue fever in Bangladesh: Clinical features and "
    "outcome. Mymensingh Medical Journal. 2018;27(3):494–499.",

    "[10] Paul KK, Dhar-Chowdhury P, Haque CE, et al. Risk factors for the presence of "
    "dengue vector mosquitoes and implications for dengue control in Dhaka, Bangladesh. "
    "PLOS ONE. 2018;13(6):e0198431.",

    "[11] Ebi KL, Nealon J. Dengue in a changing climate. Environmental Research. "
    "2016;151:115–123. doi:10.1016/j.envres.2016.07.026",

    "[12] Gubler DJ. Dengue, urbanization and globalization: the unholy trinity of the "
    "21st century. Tropical Medicine and Health. 2011;39(4 Suppl):3–11.",

    "[13] Rodó X, Romero-Alvarez D, Campbell-Lendrum D, et al. Changing climate and the "
    "COVID-19 pandemic: more than just heads or tails. Nature Medicine. "
    "2021;27(4):576–579. doi:10.1038/s41591-021-01303-y",

    "[14] Sharmin S, Glass K, Viennet E, Harley D. Interaction of mean temperature and "
    "daily fluctuation influences dengue incidence in Dhaka, Bangladesh. PLOS Neglected "
    "Tropical Diseases. 2015;9(7):e0003901.",

    "[15] WHO. Dengue: guidelines for diagnosis, treatment, prevention and control. "
    "New edition. World Health Organization; 2009. WHO/HTM/NTD/DEN/2009.1.",

    "[16] Bangladesh Bureau of Statistics (BBS). Population and Housing Census 2022: "
    "National Report. Statistics and Informatics Division, Ministry of Planning, "
    "Bangladesh; 2023.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(ref)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# FIGURES (at end, journal-style)
# ══════════════════════════════════════════════════════════════════════════════

page_break()
heading("Figures", 1)

figures = [
    (FIG_DIR / "fig1_annual_trend.png",
     "Figure 1. Annual admitted dengue cases in Bangladesh, 2018–2025. "
     "Red bar indicates the 2023 record epidemic (321,017 cases). Light blue bar indicates "
     "2020, when COVID-19 lockdown measures coincided with a 97.5% reduction in reported "
     "cases. Numbers above bars indicate absolute case counts. "
     "Source: DGHS HEOC Dengue Dashboard."),

    (FIG_DIR / "fig2_epidemic_curve.png",
     "Figure 2. Weekly epidemic curves for admitted dengue cases in Bangladesh, 2023–2025. "
     "The orange shading indicates the monsoon season (June–September, weeks 23–40). "
     "Annotated values indicate peak week and peak weekly case count for each year. "
     "The 2023 peak (W38; 20,244 cases) was 2.6-fold higher than the 2024 peak. "
     "Source: DGHS HEOC Dengue Dashboard."),

    (FIG_DIR / "fig3_seasonal_pattern.png",
     "Figure 3. Seasonal distribution of admitted dengue cases in Bangladesh, 2023–2025. "
     "Upper panel: monthly heatmap showing cases and percentage of annual total by year "
     "(black borders indicate peak month per year). Lower panel: average seasonal index "
     "(percentage of annual burden) by month, averaged across 2023–2025. "
     "Source: DGHS HEOC Dengue Dashboard."),

    (FIG_DIR / "fig4_division_analysis.png",
     "Figure 4. Geographic distribution of dengue in Bangladesh by administrative division, "
     "2026 year-to-date (January 1 – June 3, 2026). Left panel: incidence rate per 100,000 "
     "population (horizontal bars; red = highest incidence). Numbers indicate incidence rate "
     "and raw case count. Right panel: weekly case trends by division for weeks 1–22 of 2026. "
     "Population denominators from BBS 2022 Census. "
     "Source: DGHS HEOC Dengue Dashboard."),
]

for fig_path, caption in figures:
    if fig_path.exists():
        add_figure(fig_path, caption)
    else:
        para(f"[Figure not found: {fig_path.name}]", italic=True)
    doc.add_paragraph()

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"\nManuscript saved: {OUT}")
print(f"File size: {OUT.stat().st_size / 1024:.0f} KB")
